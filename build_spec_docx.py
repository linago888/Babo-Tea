import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.path.append(
    r"C:\Users\ASUS\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa


OUT = "Global Boba Graph 網站詳細功能規格書.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E9E2D7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    content_width = section_content_width_dxa(document.sections[-1])
    column_widths = column_widths_from_weights(widths or [1] * len(headers), content_width)
    apply_table_geometry(table, column_widths, table_width_dxa=content_width, indent_dxa=0)
    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    return paragraph


def add_para(document, text):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
styles["Normal"].font.size = Pt(10.5)

for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(size)
    style.font.bold = True

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Global Boba Graph\n網站詳細功能規格書")
run.font.color.rgb = RGBColor(32, 32, 32)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("全球珍珠奶茶知識與產業媒體平台 MVP Prototype 規劃").italic = True

meta = doc.add_table(rows=5, cols=2)
meta.style = "Table Grid"
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for left, right, row in [
    ("文件目的", "將網站功能、資訊架構、資料模型與 admin 後台規劃整理成可執行規格", 0),
    ("產品定位", "全球珍珠奶茶品牌、城市、飲品、新聞與趨勢的結構化資料平台", 1),
    ("適用階段", "MVP + Phase 5 已上線；Phase 5C（多帳號）與後續 3-12 個月成長", 2),
    ("文件版本", "v1.3（對齊 Phase 5A/5B/5D 實作）", 3),
    ("輸出日期", "2026-05-30", 4),
]:
    set_cell_text(meta.rows[row].cells[0], left, bold=True)
    set_cell_text(meta.rows[row].cells[1], right)
    set_cell_shading(meta.rows[row].cells[0], "E9E2D7")
meta_width = section_content_width_dxa(doc.sections[-1])
apply_table_geometry(
    meta,
    column_widths_from_weights([1.2, 4.8], meta_width),
    table_width_dxa=meta_width,
    indent_dxa=0,
)

doc.add_page_break()

add_heading(doc, "1. 產品定位", 1)
add_para(doc, "Global Boba Graph 是一個中文為主、可國際化擴充的全球珍珠奶茶知識與產業媒體平台。第一階段重點不是一般部落格，而是將品牌、城市、飲品、新聞與趨勢建立成可搜尋、可關聯、可 SEO 索引的結構化資料系統。")
add_bullets(doc, [
    "幫一般讀者快速理解全球珍奶品牌、城市與飲品趨勢。",
    "幫產業讀者追蹤品牌擴張、新品、加盟與市場變化。",
    "透過結構化資料與內部連結累積長期 SEO 流量。",
    "為未來 AI 搜尋、推薦、趨勢分析與市場情報產品建立資料基礎。",
])
add_para(doc, "文件分工：本檔聚焦在使用者體驗、頁面結構與功能模組；完整資料庫 schema、欄位、enum、關聯表與衍生指標公式請見 data-model.md。任一處的欄位定義不一致時，以 data-model.md 為準。")

add_heading(doc, "2. 使用者角色與需求", 1)
add_table(doc, ["角色", "核心需求", "對應功能"], [
    ["一般消費者", "找城市珍奶、理解飲品、比較品牌", "城市指南、飲品百科、品牌頁"],
    ["產業觀察者", "追蹤品牌擴張、新品與市場趨勢", "新聞、趨勢、品牌關聯"],
    ["加盟商 / 投資者", "了解品牌規模、市場熱度與展店區域", "品牌資料庫、城市成熟度、趨勢分數"],
    ["編輯 / 營運", "快速發布結構化內容", "CMS、資料欄位、關聯標籤"],
    ["企業客戶", "取得市場情報或 API", "報告、儀表板、資料 API"],
], widths=[1.35, 2.25, 2.55])

add_heading(doc, "3. MVP 網站導覽架構", 1)
add_para(doc, "第一階段主導覽建議如下：")
add_bullets(doc, ["首頁", "品牌", "城市", "飲品", "新聞", "趨勢", "搜尋", "訂閱"])
add_table(doc, ["頁面", "建議 URL", "主要目的"], [
    ["首頁", "/", "資料平台入口、搜尋、最新內容與核心模組總覽"],
    ["品牌列表", "/brands", "搜尋與篩選全球珍奶品牌"],
    ["品牌詳情", "/brands/gong-cha", "品牌故事、基本資料、相關城市、招牌飲品與新聞"],
    ["城市詳情", "/cities/tokyo", "城市珍奶指南、平均價格、熱門品牌與本地洞察"],
    ["飲品詳情", "/drinks/brown-sugar-milk-tea", "飲品百科、熱量、咖啡因、代表品牌與趨勢"],
    ["新聞詳情", "/news/chagee-expands-in-la", "產業新聞、來源、摘要與關聯資料"],
    ["搜尋", "/search", "跨品牌、城市、飲品與新聞的搜尋入口"],
], widths=[1.2, 2.1, 3.0])

add_heading(doc, "4. 首頁功能規格", 1)
add_table(doc, ["模組", "功能說明", "MVP 優先級"], [
    ["主視覺與定位", "清楚傳達平台是全球珍奶知識與產業資料庫，而非一般部落格", "高"],
    ["全站搜尋", "可搜尋品牌、城市、飲品與新聞", "高"],
    ["平台指標", "顯示品牌數、城市數、飲品數、新聞數與索引頁面數", "中"],
    ["最新新聞", "展示產業更新、品牌擴張與新品資訊", "高"],
    ["精選品牌", "展示重點品牌與其招牌飲品、主要市場", "高"],
    ["熱門城市", "展示城市指南、平均價格與市場成熟度", "高"],
    ["趨勢飲品", "展示熱門飲品、趨勢分數與關聯品牌", "中"],
    ["知識圖譜預覽", "顯示品牌、城市、飲品、新聞之間的關聯", "中"],
    ["訂閱 CTA", "收集電子報訂閱名單", "中"],
], widths=[1.45, 3.9, 0.85])

add_heading(doc, "5. 核心功能模組", 1)
add_para(doc, "本章描述各模組的前台呈現與資料來源；完整欄位、enum 與關聯表請見 data-model.md 對應章節。")

add_heading(doc, "5.1 品牌模組", 2)
add_para(doc, "品牌頁是 SEO 與資料價值核心，應可支撐品牌比較、產業分析與未來加盟情報。對應 data-model.md §1.1 brands、§3 companies / brand_company、§8.1 brand_drinks、§8.2 brand_cities、§8.5 brand_similarities。")
add_table(doc, ["功能", "內容"], [
    ["品牌列表", "品牌名稱、國家、成立年份、門店數量（由 stores 聚合）、主要市場、招牌飲品、品牌定位、趨勢分數（衍生）"],
    ["篩選條件", "國家、品牌定位（taxonomy）、展店狀態（brand_cities.status）、價格帶（price_tier enum）、經營模式（business_model: direct | franchise | hybrid | licensed）"],
    ["品牌詳情", "品牌摘要、品牌故事、基本資料、母公司鏈（brand_company）、全球市場分布、招牌飲品、相關城市、相關新聞、相似品牌、SEO FAQ（多語）"],
    ["關鍵欄位", "slug、name_i18n、country_code、founded_year、headquarters_city_id、business_model、price_tier、positioning_tags、social_handles、description_i18n、seo_i18n、claimed_by_user_id、verified"],
], widths=[1.35, 4.8])

add_heading(doc, "5.2 城市模組", 2)
add_para(doc, "對應 data-model.md §1.2 cities、§2 stores、§8.2 brand_cities、§8.4 drink_cities。market_maturity 為衍生指標，計算見 §7.2。")
add_table(doc, ["功能", "內容"], [
    ["城市列表", "城市、國家、平均價格（avg_price_local + currency）、市場成熟度（衍生）、熱門品牌、熱門飲品、最新新聞"],
    ["城市詳情", "城市珍奶市場概覽、推薦品牌與店鋪（stores，可上地圖）、平均價格區間、熱門飲品（含季節權重）、本地文化洞察、近期新聞、相關品牌、SEO FAQ"],
    ["關鍵欄位", "slug、name_i18n、country_code、admin_region、lat、lng、timezone、population、avg_price_local、avg_price_currency、market_maturity、description_i18n、seo_i18n"],
    ["SEO 題型", "依 locale 自動切換：東京珍奶指南：熱門品牌、價格與在地趨勢；Best Bubble Tea in Tokyo: Brands, Prices and Local Trends；東京タピオカガイド：人気ブランド・価格・最新トレンド"],
], widths=[1.35, 4.8])

add_heading(doc, "5.3 飲品模組", 2)
add_para(doc, "對應 data-model.md §1.3 drinks，飲品屬性走 §4 taxonomies 受控詞彙（tea_base / milk_type / topping / sweetener / flavor_tag / positioning_tag）。")
add_table(doc, ["功能", "內容"], [
    ["飲品列表", "飲品名稱、類型（category）、茶底、典型甜度、熱量區間、咖啡因區間、趨勢分數（衍生）、代表品牌"],
    ["篩選條件", "茶底、奶種、配料（多選）、是否含咖啡因、溫度、品類"],
    ["飲品詳情", "飲品介紹、主要成分（拆 tea_base / milk_type / toppings / sweetener，皆查 taxonomies 取多語標籤）、口味輪廓（雷達圖：sweet / bitter / milky / fruity / floral / roasted 各 0-5）、常見變體、熱量與咖啡因、代表品牌（含 local_name_i18n 與 price_local）、熱門城市、相關新聞"],
    ["關鍵欄位", "slug、name_i18n、category、tea_base[]、milk_type、toppings[]、sweetener、temperature[]、typical_sugar_levels[]、calories_kcal_min/max、caffeine_mg_min/max、flavor_profile jsonb"],
], widths=[1.35, 4.8])

add_heading(doc, "5.4 新聞模組", 2)
add_para(doc, "新聞不是孤立文章，必須關聯品牌、城市、飲品與趨勢。對應 data-model.md §1.4 news、§5 sources、§8.3 news_brands / news_cities / news_drinks。")
add_table(doc, ["新聞分類（code）", "對外標籤", "說明"], [
    ["expansion", "品牌擴張", "海外展店、加盟、通路與品牌策略"],
    ["launch", "新品上市", "新飲品、季節限定、聯名商品"],
    ["franchise-investment", "加盟與投資", "融資、加盟招商、供應鏈投資"],
    ["city-market", "城市市場", "特定城市的消費變化、商圈、價格與品牌密度"],
    ["trend", "消費趨勢", "低糖、健康化、抹茶、奶蓋、水果茶等趨勢"],
    ["supply-chain", "供應鏈", "茶葉、珍珠、奶源、包材與設備"],
    ["culture", "文化與生活", "社群流行、地方文化、消費者行為"],
], widths=[1.65, 1.4, 3.1])

add_heading(doc, "5.5 AI 摘要審核流", 2)
add_para(doc, "新聞 news.ai_summary_i18n 由 ingest pipeline 產生；ai_summary_reviewed_by 為 NULL 時前台不顯示。編輯在 CMS 看到「待審 AI 摘要」清單，確認或修改後設 ai_summary_reviewed_by = current_user，前台才顯示「AI 摘要（已編輯審核）」區塊。公信力要求：未審 AI 摘要絕對不對外公開。")

add_heading(doc, "6. 搜尋與篩選", 1)
add_table(doc, ["階段", "功能"], [
    ["MVP", "關鍵字搜尋、類型篩選、即時結果更新"],
    ["第二階段", "Algolia、自動完成、同義詞、多語搜尋、搜尋分析"],
], widths=[1.3, 4.85])

add_heading(doc, "7. 知識圖譜與關聯資料", 1)
add_para(doc, "MVP 不需要立即導入完整 graph database，但資料關聯必須先設計好，讓每個頁面都能形成內部連結與推薦。完整 schema、欄位、PK 與索引以 data-model.md §8 為準。")
add_table(doc, ["語意關聯", "對應表（data-model §8）", "關鍵欄位"], [
    ["品牌擁有飲品", "brand_drinks", "is_signature、local_name_i18n、price_local"],
    ["品牌進入城市", "brand_cities", "entered_at、store_count_cached、status"],
    ["品牌有實體店鋪", "stores（§2，新增）", "lat / lng、is_flagship、opened_at、franchise"],
    ["品牌屬於集團", "brand_company（§3，新增）", "relation、since / until"],
    ["城市流行飲品", "drink_cities", "popularity_score、seasonality"],
    ["新聞提及品牌 / 城市 / 飲品", "news_brands / news_cities / news_drinks", "relevance（primary | secondary | mentioned）、auto_tagged"],
    ["飲品形成趨勢", "metrics_daily（§7，新增）", "每日 trending_score 歷史"],
    ["相似品牌", "brand_similarities（§8.5，新增）", "score、factors"],
], widths=[1.5, 2.3, 2.35])
add_para(doc, "所有關聯表的 relevance / relation / factors 欄位日後可直接映射為 graph edge property，v2 遷移到 Neo4j / DGraph 不需重設計。")

add_heading(doc, "7.1 衍生指標（derived metrics）", 2)
add_para(doc, "trending_score、market_maturity、popularity_score 等分數不直接寫回主實體欄位，改存 metrics_daily 歷史表並可重算可審計。完整公式、權重、缺值處理見 data-model.md §7。")
add_table(doc, ["指標", "適用實體", "計算摘要"], [
    ["trending_score（0-100）", "brand / city / drink", "新聞量、新增店數、社群提及、搜尋量做 z-score 加權後套 sigmoid；加上 30 日成長率為 delta_factor。"],
    ["market_maturity", "city（分桶 emerging / growing / mature / saturated）", "活躍店數、品牌多樣性、近 24 個月新增店數加權後以全城市百分位分桶，每月 rebalance。"],
    ["popularity_score", "drink × city", "由 brand_drinks × stores × 銷售訊號（新聞、社群）計算，含季節權重。"],
], widths=[2.1, 1.7, 2.35])
add_para(doc, "呈現原則：前台只顯示最新值，可加「過去 8 週迷你曲線」；後台可展開 metrics_daily.inputs jsonb 給出原始輸入解釋；公式調整不必回填欄位，重跑 cron 即可。")

add_heading(doc, "8. 後台架構（Phase 4-5 已上線）", 1)
add_heading(doc, "8.1 入口與認證", 2)
add_bullets(doc, [
    "路徑：/admin（不含 locale prefix，與公開站 /[locale]/... 平行）",
    "認證：HTTP Basic Auth，env vars ADMIN_USER / ADMIN_PASSWORD，在 proxy.ts middleware 攔截",
    "多帳號 / RBAC：Phase 5C 規劃（v1 仍是單一帳號）",
    "Admin UI 語言：依 NEXT_LOCALE cookie（公開站切換語言會跟著切）",
    "響應式：≥ 640px 顯示左側固定 sidebar；< 640px 改用上方漢堡按鈕 + 抽屜選單",
])

add_heading(doc, "8.2 已上線的 admin 模組", 2)
add_table(doc, ["模組", "路徑", "功能摘要"], [
    ["總覽", "/admin", "8 個磚卡（即時筆數）+ 各 Quality 入口"],
    ["品牌", "/admin/brands", "list + 篩選 + search；edit 7 個 tabs（基本/i18n/SEO/進階/招牌飲品/進駐城市/母公司）"],
    ["城市", "/admin/cities", "list + 篩選 + search；edit（含 lat/lng/timezone/avg_price/market_maturity）"],
    ["飲品", "/admin/drinks", "list + category/status 篩選；edit（含 recipe + nutrition + flavor_profile JSON）"],
    ["新聞", "/admin/news", "list + 篩選；edit 7 個 tabs（基本/i18n本文/SEO/AI摘要/相關品牌/城市/飲品）"],
    ["來源", "/admin/sources", "list + search；slug 與 domain 雙 unique"],
    ["詞彙", "/admin/taxonomies", "按 kind 分組；(kind, code) compound unique；parent 階層"],
    ["公司", "/admin/companies", "list + search；slug / country / ticker / website"],
    ["門市", "/admin/stores", "brand + city FK + lat/lng + opening_hours JSON + isFlagship / franchise"],
    ["內容品質", "/admin/quality", "completeness 分佈、orphan、AI 待審、review_due 清單"],
    ["指標趨勢", "/admin/metrics", "metrics_daily 視覺化（Top 10 × 4 排行榜 + sparkline + Δ）"],
    ["搜尋紀錄", "/admin/search-log", "search_log 分析（熱門 / 零結果 / 日趨勢 / 語系與國家分佈）"],
], widths=[1.0, 1.7, 3.45])

add_heading(doc, "8.3 編輯表單共通設計", 2)
add_bullets(doc, [
    "多語 i18n tab：每個 i18n 欄位（name / description / SEO）都有 4 個 locale 子分頁",
    "儲存後 SSG 自動失效：API route 呼叫 revalidatePath('/[locale]/{entity}', 'layout')",
    "完整度即時重算：cities / drinks / news 儲存時 scoreXxx() 立刻重算寫回 DB",
    "軟刪除：「刪除」實為 status = ARCHIVED，公開站隱藏，後台仍可看",
    "Zod 422 / 409 / 500 錯誤顯示在表單頂",
])

add_heading(doc, "8.4 關聯資料編輯（Phase 5A）", 2)
add_para(doc, "Brand 編輯頁多 3 個 tab：招牌飲品（brand_drinks）、進駐城市（brand_cities）、母公司（brand_company）。News 編輯頁多 3 個 tab：相關品牌 / 城市 / 飲品（news_brands / news_cities / news_drinks）。")
add_para(doc, "寫入策略：主表 + 關聯改動包在 prisma.$transaction 內，採 delete-all + insert-all — 簡單、可靠、不留 orphan。對 brand_company（複合 PK 含 since）特別重要。")

add_heading(doc, "8.5 AI 草稿生成（Phase 5B）", 2)
add_para(doc, "每個編輯頁的對應 tab 多了「✨ AI 補完」按鈕。按一次按鈕，4 個 locale 同時填好。Endpoint：POST /api/admin/ai/draft；模型：OpenAI gpt-4o-mini；透過 Vercel AI SDK generateObject + Zod schema 確保輸出形狀正確。")
add_table(doc, ["模組", "AI 補完位置"], [
    ["Brand", "i18n description + SEO (title+desc)"],
    ["City", "i18n description + SEO"],
    ["Drink", "i18n description + SEO"],
    ["News", "i18n summary + body（完整 markdown 文章）+ SEO + advanced AI 摘要"],
], widths=[1.2, 4.95])
add_para(doc, "需設 Vercel env var OPENAI_API_KEY。未設時按鈕仍顯示，按下去跳 503 提示。系統 prompt 明確指示「不要編造數字、引言、店數」，只用 form 提供的 context 作生成依據。")

add_heading(doc, "8.6 圖片上傳（Phase 5B）", 2)
add_bullets(doc, [
    "Brand 的 logoUrl 與 News 的 heroImageUrl 改成 URL + 拖檔上傳 + 預覽三合一元件",
    "URL 文字欄向後相容（編輯仍可貼外部連結）",
    "📁 選檔按鈕 + 拖放區（直接拖圖到框內）",
    "即時預覽縮圖",
    "上傳到 Vercel Blob，回傳 public URL 自動填回",
    "需設 Vercel env var BLOB_READ_WRITE_TOKEN（Vercel Storage → Create Blob Store 自動注入）",
    "限制：≤ 5 MB；MIME whitelist：jpeg/png/webp/gif/svg/avif",
])

add_heading(doc, "8.7 發布前 gating 檢查", 2)
add_bullets(doc, [
    "由 src/lib/content-quality/gating.ts 的 Zod schema 強制",
    "slug 唯一性（DB 層 @unique + API 409 衝突檢查）",
    "seo_i18n[預設 locale] 的 title 與 description 已填",
    "至少一個關聯實體（news_brands ∪ news_cities ∪ news_drinks）",
    "summary_i18n 已填",
    "source_id 設定（新聞）",
    "hero image 與 alt text",
    "JSON-LD 必填欄位齊全（依模板對應 schema.org type）",
])

add_heading(doc, "9. 分析儀表板（Phase 5D）", 1)

add_heading(doc, "9.1 內容品質儀表板（/admin/quality）", 2)
add_bullets(doc, [
    "各實體 completeness_score 分佈（≥80 / 50-79 / <50 / 未評分 4 桶 × 4 entity = 16 個磚卡）",
    "最低完整度前 5 名（每個 entity）",
    "orphan 清單（無任何關聯的孤兒資料）",
    "AI 摘要待審清單（ai_summary_reviewed_at IS NULL）",
    "review_due_at 已到期的實體",
    "點任一筆名稱直接跳到 admin 編輯頁",
])

add_heading(doc, "9.2 指標趨勢儀表板（/admin/metrics）", 2)
add_para(doc, "用既有的 metrics_daily 表，4 個 Top 10 排行榜：熱度最高品牌（trending_score）、近 30 天新聞最多品牌（news_count_30d）、城市熱度榜（popularity_score）、飲品流行度榜（popularity_score）。")
add_para(doc, "每一行顯示：名次 + 名稱（連 admin 編輯頁）+ 30 天 SVG sparkline + 當前值 + Δ vs 7 天前（↑↓ 帶顏色）。查詢策略：DISTINCT ON (entity_id) 拿每實體最新值 → top N entity_ids → IN-query 拉 30 天序列。")

add_heading(doc, "9.3 搜尋紀錄儀表板（/admin/search-log）", 2)
add_bullets(doc, [
    "4 個磚卡：7 天 / 30 天搜尋總數、7 天零結果次數、活躍語系數",
    "30 天 daily volume sparkline（含空白日對齊）",
    "熱門關鍵字 Top 20（近 7 天）：query × count × 平均命中數",
    "零結果關鍵字 Top 20（近 30 天）：橘色 bar，編輯團隊看完就知道下一篇要寫什麼",
    "語系分佈 + 國家分佈長條圖",
])
add_para(doc, "對應 data-model.md §14 search_log 表。寫入策略：/[locale]/search 渲染後 void logSearch(...)，fire-and-forget 不阻塞使用者。< 2 字元 query 不記。不記 IP，只記 country code。")

add_heading(doc, "10. Prototype 內容（已凍結，搬到 prototype/）", 1)
add_para(doc, "歷史靜態互動原型保留作早期構思參考，Phase 1 之後已用 Next.js + Tailwind 4 + i18n 重做整套，prototype/ 不再更新。")

add_heading(doc, "11. 開發進度（已完成回顧）", 1)
add_table(doc, ["Phase", "內容"], [
    ["1 Day 1-2", "Next.js 16 + TS + Tailwind 4 + Prisma 7 + next-intl scaffold；Supabase 連線；Vercel 部署"],
    ["1 Day 3-6", "6 張主表 + 7 張關聯表 + stores / companies + metrics_daily"],
    ["1.5", "seed data + idempotent upsert + metrics 首次計算"],
    ["2.1-2.7", "公開站 5 大模組（首頁 / brands / cities / drinks / news）+ BrandLogo"],
    ["2.8", "跨實體 search + /search 頁"],
    ["3", "sitemap 套件 × 4 + News sitemap + robots + RSS + 動態 OG image"],
    ["4", "/admin HTTP Basic Auth + /admin/quality 內容品質儀表板"],
    ["4.5-4.6", "admin i18n + 全 8 個 entity CRUD（brands / cities / drinks / news / sources / taxonomies / companies / stores）"],
    ["5A", "Brand/News 關聯編輯 tabs（brand_drinks / brand_cities / brand_company / news_*）"],
    ["5B", "AI 草稿生成（OpenAI gpt-4o-mini）+ Vercel Blob 圖片上傳"],
    ["5C", "🚧 NextAuth 多帳號 + RBAC + audit log（待規劃）"],
    ["5D", "Metrics dashboard + Search log 表與儀表板"],
], widths=[1.0, 5.15])

add_heading(doc, "12. 後續成長功能", 1)
add_table(doc, ["階段", "功能"], [
    ["待規劃", "Phase 5C 多帳號 + RBAC + audit log；進階搜尋（Meilisearch / Typesense 自架）；每週電子報實作；城市市場比較頁（/compare?a=tokyo&b=taipei）；品牌認領（claimed_by_user_id / verified）；News crawler / RSS aggregator + AI 自動 tag"],
    ["3-12 個月", "全球珍奶地圖（stores × cities.lat/lng，Leaflet / Mapbox）、AI 問答助手（以結構化資料為 RAG 來源）、加盟商情報（brand_company × stores × metrics_daily）、市場報告（自動產生 monthly PDF）、企業資料 API（slug 與 UUID 雙鍵已備，加 API key + rate limit）"],
], widths=[1.3, 4.85])

add_heading(doc, "13. 建議結論", 1)
add_para(doc, "Global Boba Graph 已從 SEO 內容資料庫起步完成 Phase 1-5（除 5C 多帳號）— 公開站、SEO 套件、完整 admin CRUD、AI 編輯輔助、Metrics 視覺化都已上線。資料模型 v1.3 仍預留 graph DB、品牌認領、付費 tier、多幣別、企業 API 等成長鈎子。下一階段重點為：1) 多帳號 / RBAC 支撐團隊編輯；2) 透過 search log 與 metrics dashboard 看出的內容缺口持續水平擴量；3) 評估全球珍奶地圖、AI 問答助手等 v2 功能的市場時機。")

doc.save(OUT)
print(OUT)
